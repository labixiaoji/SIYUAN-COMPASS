from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/思源Compass测评问卷清单.docx")


SECTIONS = [
    (
        "一、基本信息",
        [
            ("你的年级是？", "单选", ["大一", "大二", "大三", "大四", "研一", "研二", "研三", "博士", "其他"], ""),
            ("你的学院和专业是？", "文本输入", [], "____________________________________________"),
            ("你的家乡或主要成长地是？", "文本输入", [], "____________________________________________"),
            (
                "你目前更倾向毕业后在哪里发展？",
                "单选",
                ["上海", "北京 / 深圳 / 广州", "长三角其他城市", "家乡或周边城市", "海外 / 境外", "暂不确定"],
                "",
            ),
        ],
    ),
    (
        "二、当前困惑",
        [
            (
                "你现在最想解决的生涯问题是什么？",
                "多选，最多选择 3 项",
                [
                    "不知道未来适合做什么",
                    "纠结就业、读研、出国、读博",
                    "不确定自己适合哪个行业",
                    "不知道本专业未来有哪些出路",
                    "想提高求职 / 实习竞争力",
                    "对未来收入、城市、生活状态感到焦虑",
                    "家庭期待和个人想法不一致",
                    "想要更清楚地规划未来5—10年",
                    "其他",
                ],
                "",
            ),
            ("如果只能用一句话描述你现在的困惑，你会怎么说？", "文本输入", [], "____________________________________________________________\n____________________________________________________________"),
        ],
    ),
    (
        "三、教育路径",
        [
            (
                "你目前毕业后的初步选择是？",
                "单选",
                ["直接就业", "国内读研", "出国 / 境外深造", "读博 / 科研方向", "考公 / 选调 / 事业单位", "国企 / 央企 / 研究院", "创业 / 自由职业", "暂不确定"],
                "",
            ),
            (
                "你选择或倾向这条路径的主要原因是什么？",
                "多选",
                ["提升学历和平台", "目标职业需要更高学历", "真正喜欢科研或学术探索", "想延缓就业压力", "家庭建议或外部期待", "想获得更稳定的发展路径", "想追求更高收入", "想换城市或拓宽机会", "还没有想清楚，只是暂时倾向", "其他"],
                "",
            ),
            ("你对这个选择的确定程度是多少？", "1 到 5 分", ["1", "2", "3", "4", "5"], ""),
            ("你是否考虑过读博？", "单选", ["明确考虑", "有一点兴趣", "不太考虑", "完全不考虑", "还不了解"], ""),
            ("如果你考虑读博，你考虑读博的主要原因是什么？", "条件题，多选", ["科研兴趣", "未来想进高校 / 科研院所", "提升专业壁垒", "暂时不想就业", "家庭或导师建议", "其他"], ""),
        ],
    ),
    (
        "四、未来愿景",
        [
            (
                "未来 5 年，你最希望优先获得什么？",
                "多选，最多选择 3 项",
                ["高收入 / 经济回报", "稳定 / 安全感", "快速成长 / 学习新技能", "专业深度 / 技术壁垒", "意义感 / 帮助他人或创造价值", "自由 / 弹性 / 不被控制", "声望 / 被认可和尊重", "平衡 / 有时间给家庭和生活", "高质量人际关系", "健康和可持续的生活状态"],
                "",
            ),
            (
                "未来 5 年，你希望主要深耕的行业或方向是？",
                "多选，最多选择 2 项",
                ["互联网 / 信息技术 / 人工智能", "高端制造 / 汽车 / 能源", "金融 / 咨询 / 商业分析", "医疗 / 生命科学 / 药企", "教育 / 文化 / 公共服务", "体制内 / 选调 / 公务员 / 事业单位", "科研 / 高校 / 学术道路", "创业 / 自由职业", "暂不确定", "其他"],
                "",
            ),
            ("未来 5—10 年，你更希望成为哪类人？", "单选", ["专业技术型人才", "科研学术型人才", "管理者 / 组织领导者", "商业分析 / 金融咨询类人才", "稳定体制内 / 国企发展者", "产品 / 运营 / 综合型人才", "创业者 / 自由职业者", "教育、公益、公共服务方向的人", "还不确定"], ""),
            ("未来 5—10 年，你理想中的生活状态更接近哪一种？", "单选", ["事业优先，愿意承受较高强度", "稳定体面，风险可控", "工作生活平衡，有时间给家人和兴趣", "高自由度，不喜欢被固定结构束缚", "保持探索，愿意接受变化和试错", "暂时不确定"], ""),
            ("请用几句话描述：10 年后，你希望自己成为怎样的人？", "文本输入", [], "____________________________________________________________\n____________________________________________________________\n____________________________________________________________"),
        ],
    ),
    (
        "五、价值、能力与兴趣",
        [
            ("请选出你最看重的 3 项价值观。", "多选，选择 3 项", ["高收入 / 经济回报", "稳定 / 安全感", "成长 / 学习新技能的机会", "意义感 / 帮助他人或创造价值", "人际关系", "自由 / 弹性 / 不被控制", "声望 / 被认可和尊重", "平衡 / 有足够时间给家庭和生活"], ""),
            ("能力自评：1 = 很不符合，5 = 非常符合。", "量表题", ["数学、逻辑推理", "写作、表达、讲故事", "空间、方向、设计", "识人、沟通、理解情绪"], ""),
            ("兴趣倾向：1 = 很不喜欢，5 = 非常喜欢。", "量表题", ["动手操作、修理工具", "研究问题、分析数据", "创作、写作、设计或表达", "帮助他人、教学或咨询", "销售、领导或影响他人", "按规则整理信息、处理细节"], ""),
        ],
    ),
    (
        "六、行动基础",
        [
            ("为了未来目标，你已经做过哪些准备？", "多选", ["修读相关课程", "参加科研 / 课题", "做过项目作品", "参加竞赛", "有实习经历", "参加社团 / 学生工作", "和老师、学长学姐或校友聊过", "参加过宣讲会 / 招聘会", "看过目标岗位招聘要求", "准备过简历", "投递过实习 / 工作", "还没有明显准备", "其他"], ""),
            ("你目前最缺的是什么？", "多选，最多选择 3 项", ["不知道方向", "不知道真实岗位要求", "缺技能", "缺项目 / 实习经历", "缺人脉和信息", "缺行动力", "缺信心", "缺时间", "缺对行业的了解", "缺清晰计划"], ""),
            ("你是否了解本专业近几届毕业生的主要去向？", "单选", ["比较了解", "听说过一些", "基本不了解", "完全不了解"], ""),
            ("你是否了解自己心仪岗位的要求？", "单选", ["比较了解，并看过岗位 JD", "有一点了解", "只知道大概方向", "基本不了解", "还没有心仪岗位"], ""),
            ("你的身体健康和精力状态如何？", "单选", ["很好，每周规律运动", "一般，偶尔运动", "较差，睡眠、压力或身体状态影响学习生活", "不想回答"], ""),
            ("你每周大概运动几次？每次多长时间？", "文本输入", [], "____________________________________________"),
        ],
    ),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, size=10.5, bold=False, color=None):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def add_options(cell, qtype, options, answer_line):
    if answer_line:
        for idx, line in enumerate(answer_line.split("\n")):
            p = cell.add_paragraph() if idx else cell.paragraphs[0]
            add_text(p, line, 10.5, False)
        return

    if qtype == "量表题":
        table = cell.add_table(rows=1, cols=6)
        table.autofit = False
        headers = ["项目", "1", "2", "3", "4", "5"]
        for idx, header in enumerate(headers):
            p = table.rows[0].cells[idx].paragraphs[0]
            add_text(p, header, 9.5, True, "1F4D78")
            set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
            set_cell_margins(table.rows[0].cells[idx])
        for option in options:
            row = table.add_row().cells
            add_text(row[0].paragraphs[0], option, 9.5)
            for i in range(1, 6):
                p = row[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text(p, "□", 10.5)
            for c in row:
                set_cell_margins(c)
        return

    prefix = "□ "
    if "单选" in qtype:
        prefix = "○ "
    for idx, option in enumerate(options):
        p = cell.add_paragraph() if idx else cell.paragraphs[0]
        add_text(p, f"{prefix}{option}", 10.5)


def build_doc():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(title, "思源 Compass 测评问卷清单", 20, True, "0F766E")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(subtitle, "用于职业规划智能体的首次填写与报告生成", 10.5, False, "555555")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(meta, "题目总数：25  |  题型：单选、多选、量表、文本输入", 9.5, False, "555555")

    q_no = 1
    for section_title, questions in SECTIONS:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        add_text(heading, section_title, 14, True, "1F4D78")

        for title_text, qtype, options, answer_line in questions:
            table = doc.add_table(rows=2, cols=3)
            table.autofit = False
            table.allow_autofit = False
            widths = [Inches(0.55), Inches(4.75), Inches(1.05)]
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    cell.width = widths[idx]
                    set_cell_margins(cell)

            header = table.rows[0].cells
            header[0].merge(header[1])
            header[0].merge(header[2])
            set_cell_shading(header[0], "E8F3F1")
            p = header[0].paragraphs[0]
            add_text(p, f"{q_no}. {title_text}", 10.5, True, "202124")
            add_text(p, f"  （{qtype}）", 9.5, False, "64635F")

            body = table.rows[1].cells
            body[0].text = ""
            p0 = body[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p0, "作答", 9.5, True, "64635F")
            body[1].merge(body[2])
            add_options(body[1], qtype, options, answer_line)

            doc.add_paragraph()
            q_no += 1

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
